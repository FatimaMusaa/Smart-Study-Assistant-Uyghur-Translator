from io import BytesIO
from typing import Literal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from pydantic import BaseModel


router = APIRouter()


class ExportDocxRequest(BaseModel):
    title: str
    original_text: str
    translated_text: str
    source_type: Literal["chapter", "page"]
    source_number: int
    review_status: str = "not_reviewed"


def add_paragraph_with_style(
    document: Document,
    text: str,
    *,
    rtl: bool = False,
    bold: bool = False,
    font_size: int = 12,
):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)

    run.bold = bold
    run.font.size = Pt(font_size)

    if rtl:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    return paragraph


@router.post("/export/docx")
async def export_docx(payload: ExportDocxRequest):
    document = Document()

    document.add_heading(payload.title, level=1)

    add_paragraph_with_style(
        document,
        f"Source Type: {payload.source_type}",
        font_size=11,
    )
    add_paragraph_with_style(
        document,
        f"Source Number: {payload.source_number}",
        font_size=11,
    )
    add_paragraph_with_style(
        document,
        f"Review Status: {payload.review_status}",
        font_size=11,
    )

    document.add_paragraph("")

    document.add_heading("Uyghur Translation", level=2)
    add_paragraph_with_style(
        document,
        payload.translated_text,
        rtl=True,
        font_size=13,
    )

    document.add_paragraph("")

    document.add_heading("Original Text", level=2)
    add_paragraph_with_style(
        document,
        payload.original_text,
        font_size=11,
    )

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    safe_filename = f"{payload.source_type}-{payload.source_number}-translation"

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{safe_filename}.docx"'
    },
)