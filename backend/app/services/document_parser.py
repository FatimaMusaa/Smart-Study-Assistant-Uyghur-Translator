import re
from typing import Any

import fitz
from docx import Document


def is_likely_real_table(rows: list[list[str]]) -> bool:
    if len(rows) < 2:
        return False

    column_counts = [len(row) for row in rows if row]

    if not column_counts:
        return False

    max_columns = max(column_counts)

    # A real table should usually have at least 2 columns.
    if max_columns < 2:
        return False

    # Avoid detecting simple bullet/word lists as tables.
    non_empty_cells = [
        cell.strip()
        for row in rows
        for cell in row
        if cell and cell.strip()
    ]

    if len(non_empty_cells) < 4:
        return False

    # If most rows only contain one useful cell, it is probably a list, not a table.
    rows_with_multiple_cells = 0

    for row in rows:
        useful_cells = [cell for cell in row if cell and cell.strip()]
        if len(useful_cells) >= 2:
            rows_with_multiple_cells += 1

    if rows_with_multiple_cells < 2:
        return False

    # If cells are mostly very short one-word list items, likely false positive.
    average_cell_length = sum(len(cell) for cell in non_empty_cells) / len(non_empty_cells)

    if average_cell_length < 2:
        return False

    return True


def extract_tables_from_page(page: fitz.Page) -> list[dict[str, Any]]:
    tables: list[dict[str, Any]] = []

    try:
        found_tables = page.find_tables()
    except Exception:
        return tables

    for table_index, table in enumerate(found_tables.tables, start=1):
        try:
            extracted_rows = table.extract()
        except Exception:
            continue

        cleaned_rows: list[list[str]] = []

        for row in extracted_rows:
            cleaned_row = []

            for cell in row:
                if cell is None:
                    cleaned_row.append("")
                else:
                    cleaned_row.append(str(cell).strip())

            if any(cell for cell in cleaned_row):
                cleaned_rows.append(cleaned_row)

        if not cleaned_rows:
            continue

        if not is_likely_real_table(cleaned_rows):
            continue

        tables.append(
            {
                "table_number": len(tables) + 1,
                "rows": cleaned_rows,
                "row_count": len(cleaned_rows),
                "column_count": max(len(row) for row in cleaned_rows),
            }
        )

    return tables

def extract_text_from_pdf(file_bytes: bytes) -> dict[str, Any]:
    doc = fitz.open(stream=file_bytes, filetype="pdf")

    pages = []
    full_text_parts = []

    for page_index, page in enumerate(doc, start=1):
        page_text = page.get_text("text").strip()
        page_tables = extract_tables_from_page(page)

        pages.append(
            {
                "page_number": page_index,
                "text": page_text,
                "tables": page_tables,
                "table_count": len(page_tables),
            }
        )

        if page_text:
            full_text_parts.append(f"--- Page {page_index} ---\n{page_text}")

    full_text = "\n\n".join(full_text_parts)
    chapters = detect_chapters_from_pages(pages)

    return {
        "text": full_text,
        "pages": pages,
        "page_count": len(pages),
        "chapters": chapters,
        "chapter_count": len(chapters),
    }


def extract_text_from_docx(file_bytes: bytes) -> dict[str, Any]:
    from io import BytesIO

    document = Document(BytesIO(file_bytes))

    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    full_text = "\n\n".join(paragraphs)

    pages = [
        {
            "page_number": 1,
            "text": full_text,
            "tables": [],
            "table_count": 0,
        }
    ]

    chapters = detect_chapters_from_pages(pages)

    return {
        "text": full_text,
        "pages": pages,
        "page_count": 1,
        "chapters": chapters,
        "chapter_count": len(chapters),
    }


def extract_text_from_file(filename: str, file_bytes: bytes) -> dict[str, Any]:
    lower_filename = filename.lower()

    if lower_filename.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)

    if lower_filename.endswith(".docx"):
        return extract_text_from_docx(file_bytes)

    raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")


def is_toc_like_line(line: str) -> bool:
    stripped = line.strip()

    if not stripped:
        return False

    if re.search(r"\.{3,}", stripped):
        return True

    if re.search(r"\s+\d+$", stripped) and len(stripped) > 12:
        return True

    return False


def page_looks_like_toc(page_text: str) -> bool:
    upper_text = page_text.upper()

    toc_keywords = [
        "TABLE OF CONTENTS",
        "CONTENTS",
        "الفهرس",
    ]

    return any(keyword in upper_text for keyword in toc_keywords)


def detect_chapters_from_pages(pages: list[dict[str, Any]]) -> list[dict[str, Any]]:
    chapter_candidates = []

    chapter_pattern = re.compile(
        r"^(CHAPTER\s+\d+(?:\s*&\s*\d+)?(?:\s*[-–]\s*.+)?|CHAPTER\s+\d+\s+.+)$",
        re.IGNORECASE,
    )

    for page in pages:
        page_number = page["page_number"]
        page_text = page.get("text", "")

        if page_looks_like_toc(page_text):
            continue

        lines = [
            line.strip()
            for line in page_text.splitlines()
            if line.strip()
        ]

        search_lines = lines[:8]

        for line in search_lines:
            if is_toc_like_line(line):
                continue

            if chapter_pattern.match(line):
                chapter_candidates.append(
                    {
                        "chapter_number": len(chapter_candidates) + 1,
                        "title": line,
                        "start_page": page_number,
                    }
                )
                break

    chapters = []

    for index, chapter in enumerate(chapter_candidates):
        start_page = chapter["start_page"]

        if index + 1 < len(chapter_candidates):
            end_page = chapter_candidates[index + 1]["start_page"] - 1
        else:
            end_page = pages[-1]["page_number"] if pages else start_page

        chapter_pages = [
            page
            for page in pages
            if start_page <= page["page_number"] <= end_page
        ]

        chapter_text_parts = []

        for page in chapter_pages:
            page_text = page.get("text", "")
            if page_text:
                chapter_text_parts.append(
                    f"--- Page {page['page_number']} ---\n{page_text}"
                )

        chapters.append(
            {
                "chapter_number": chapter["chapter_number"],
                "title": chapter["title"],
                "start_page": start_page,
                "end_page": end_page,
                "text": "\n\n".join(chapter_text_parts),
            }
        )

    return chapters